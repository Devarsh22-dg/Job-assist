import streamlit as st
from collections import Counter
import re
from datetime import date
import difflib
from io import BytesIO

# Prefer PyPDF2, fall back to pypdf if available
try:
    from PyPDF2 import PdfReader
except ModuleNotFoundError:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError:
        PdfReader = None

# docx comes from the python-docx package; handle if it's missing so the app doesn't crash
try:
    import docx
    from docx import Document
except ModuleNotFoundError:
    docx = None
    Document = None

# PDF export library (fpdf2)
try:
    from fpdf import FPDF
except ModuleNotFoundError:
    FPDF = None

# -----------------------------
# Domain-Aware Keyword Extraction (existing)
# -----------------------------
STOPWORDS = set("""
a an the and or but for from in on to with by as at of be been being am is are was were will can must
should could may might would have has had do does did not your their our my its you they we i this that
these those such other more some many few about it them us her him she he where when how what who which
if then else than because since until while although though before after above below between during per
each every own same so very into over under around out up down off across near far much also even
""".split())

def extract_jargon_keywords(text, max_count=40):
    if not text:
        return []
    phrases = re.findall(r'\b[a-zA-Z]{3,}(?:\s+[a-zA-Z]{3,}){0,2}\b', text.lower())
    filtered = [
        p.strip() for p in phrases
        if not all(w in STOPWORDS for w in p.split())
        and len(p.split()) <= 3
        and not p.isdigit()
    ]
    freq = Counter(filtered)
    ranked = sorted(freq.items(), key=lambda x: (len(x[0].split()), x[1]), reverse=True)
    keywords = [term for term, _ in ranked[:max_count]]
    return list(dict.fromkeys(keywords))

# -----------------------------
# Tailoring and cover letter (existing)
# -----------------------------
def tailor_resume(resume_text, job_text):
    keywords = extract_jargon_keywords(job_text)
    lines = [l.strip() for l in resume_text.split("\n") if l.strip()]
    tailored_lines = []
    lower_resume = " ".join(lines).lower()
    missing = [k for k in keywords if k not in lower_resume]

    for line in lines:
        matched = [k for k in keywords if k in line.lower()]
        if matched:
            tailored_lines.append(f"- {line}  (keywords: {', '.join(matched)})")
        else:
            tailored_lines.append(f"- {line}")

    suggestions = [
        f"- Suggested: Add details or examples demonstrating experience with '{kw}'."
        for kw in missing[:10]
    ]
    return "\n".join(tailored_lines + ["", "--- Suggested Additions ---", *suggestions])

def generate_cover_letter(name, company, position, job_text, summary):
    keywords = extract_jargon_keywords(job_text, 10)
    top_keywords = ", ".join(keywords[:6])
    today = date.today().strftime("%B %d, %Y")
    return (
        f"{today}\n\n"
        f"Dear Hiring Team at {company},\n\n"
        f"I am excited to apply for the {position} role. My background aligns closely with this opportunity, "
        f"especially in areas such as {top_keywords}. {summary}\n\n"
        f"I look forward to contributing to {company}'s success.\n\n"
        f"Sincerely,\n{name}"
    )

# -----------------------------
# File reading helpers
# -----------------------------
def read_file(uploaded_file):
    if uploaded_file is None:
        return "", ""
    file_type = uploaded_file.name.lower()
    if file_type.endswith(".txt"):
        try:
            txt = uploaded_file.read().decode("utf-8")
        except Exception:
            txt = uploaded_file.getvalue().decode("utf-8") if hasattr(uploaded_file, "getvalue") else ""
        return txt, ".txt"
    elif file_type.endswith(".pdf"):
        if PdfReader is None:
            st.error("PDF read support missing. Add PyPDF2 or pypdf to requirements.txt and redeploy.")
            return "", ".pdf"
        try:
            reader = PdfReader(uploaded_file)
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
            return text, ".pdf"
        except Exception as e:
            st.error(f"Failed to read PDF: {e}")
            return "", ".pdf"
    elif file_type.endswith(".docx"):
        if docx is None:
            st.error("DOCX read support missing. Add python-docx to requirements.txt and redeploy.")
            return "", ".docx"
        try:
            document = docx.Document(uploaded_file)
            text = "\n".join(p.text for p in document.paragraphs)
            return text, ".docx"
        except Exception as e:
            st.error(f"Failed to read DOCX: {e}")
            return "", ".docx"
    else:
        st.warning("Unsupported file type. Please upload .txt, .pdf or .docx")
        return "", ""

# -----------------------------
# Unified colored diff generator (word-level with replace handling)
# -----------------------------
def make_colored_unified_html(orig_text, new_text):
    """
    Build a unified HTML representation of the diff, with:
      - Additions: blue background
      - Replacements: show original struck-through red followed by new yellow-highlighted text
      - Deletions: red strike-through
    """
    # Work at word level for fine-grained diff
    orig_words = orig_text.split()
    new_words = new_text.split()
    sm = difflib.SequenceMatcher(a=orig_words, b=new_words)
    parts = []
    # HTML/CSS styles
    styles = {
        "add": "background:#d8ebff;padding:2px;border-radius:3px;",         # blue-ish
        "del": "background:#ffd6d6;text-decoration:line-through;padding:2px;border-radius:3px;",  # red strike
        "rep_new": "background:#fff7cc;padding:2px;border-radius:3px;",     # yellow
        "rep_old": "background:#ffd6d6;text-decoration:line-through;padding:2px;border-radius:3px;", # red old
        "normal": ""
    }

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            parts.append(escape_html(" ".join(orig_words[i1:i2])))
        elif tag == 'insert':
            added = " ".join(new_words[j1:j2])
            parts.append(f"<span style='{styles['add']}'>{escape_html(added)}</span>")
        elif tag == 'delete':
            deleted = " ".join(orig_words[i1:i2])
            parts.append(f"<span style='{styles['del']}'>{escape_html(deleted)}</span>")
        elif tag == 'replace':
            old = " ".join(orig_words[i1:i2])
            new = " ".join(new_words[j1:j2])
            # show old (red strike) then new (yellow)
            parts.append(f"<span style='{styles['rep_old']}'>{escape_html(old)}</span>")
            parts.append(f"<span style='{styles['rep_new']}'>{escape_html(new)}</span>")
    # Join with spaces, wrap in <div> for scrolling
    html_body = "<div style='line-height:1.6;font-family: -apple-system, BlinkMacSystemFont, \"Segoe UI\", Roboto, \"Helvetica Neue\", Arial; padding:8px;'>"
    html_body += " ".join(parts)
    html_body += "</div>"
    return html_body

def escape_html(s: str) -> str:
    # minimal HTML escaping
    return (s.replace("&", "&amp;")
             .replace("<", "&lt;")
             .replace(">", "&gt;")
             .replace('"', "&quot;"))

# -----------------------------
# Export helpers: DOCX and PDF (text-based ATS-friendly)
# -----------------------------
def export_docx_text(tailored_text: str) -> bytes:
    """Return a bytes object of a simple text-based DOCX with the tailored content."""
    if Document is None:
        raise RuntimeError("python-docx not installed")
    doc = Document()
    for line in tailored_text.splitlines():
        doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def export_pdf_text(tailored_text: str) -> bytes:
    """Return a bytes object of a simple text-based PDF with the tailored content."""
    if FPDF is None:
        raise RuntimeError("fpdf2 not installed")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    # write line by line
    for line in tailored_text.splitlines():
        # FPDF's cell/writing: wrap if needed
        pdf.multi_cell(0, 7, line)
    bio = BytesIO()
    pdf.output(bio)
    bio.seek(0)
    return bio.read()

# -----------------------------
# Streamlit UI
# -----------------------------
st.set_page_config(page_title="AI Job Application Agent", layout="wide")
st.title("🤖 AI Job Application Agent — Unified Color Diff + Download")
st.write("Unified, color-coded diff: Blue = additions, Yellow = replacement (new), Red = deletions (strike-through).")

with st.sidebar:
    st.header("Settings")
    name = st.text_input("Your Name", "Your Name")
    company = st.text_input("Company", "Company")
    position = st.text_input("Position", "Position")
    summary = st.text_area("Short Resume Summary", "I have X years of experience delivering measurable results.")
    st.markdown("---")
    st.caption("Upload a resume (txt/pdf/docx). The tailored resume can be downloaded as the same format (text-based).")

uploaded_file = st.file_uploader("📤 Upload your resume (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"])
resume_text, orig_ext = ("", "")
if uploaded_file is not None:
    resume_text, orig_ext = read_file(uploaded_file)
    if resume_text:
        st.success("Resume uploaded successfully!")

job_desc = st.text_area("📋 Paste Job Description:", height=200)

if st.button("✨ Tailor Resume"):
    if not job_desc.strip():
        st.warning("Please paste a job description first.")
    else:
        # generate
        keywords = extract_jargon_keywords(job_desc, 40)
        tailored_resume = tailor_resume(resume_text, job_desc)
        cover_letter = generate_cover_letter(name, company, position, job_desc, summary)

        st.subheader("🧩 Extracted Domain Keywords")
        st.write(", ".join(keywords) if keywords else "No jargon-like keywords detected.")

        st.subheader("🔀 Unified Color-Coded Diff (Original → Tailored)")
        diff_html = make_colored_unified_html(resume_text, tailored_resume)
        # provide a scrollable container
        st.markdown(
            "<div style='max-height:400px;overflow:auto;border:1px solid #eee;padding:8px;border-radius:6px;'>"
            + diff_html +
            "</div>",
            unsafe_allow_html=True,
        )

        st.subheader("💌 Cover Letter")
        st.text_area("Generated Cover Letter", cover_letter, height=200)

        # Download tailored resume as same format type (text-based)
        # Fallback: if original extension unknown, default to txt
        if orig_ext == ".docx":
            try:
                docx_bytes = export_docx_text(tailored_resume)
                st.download_button("⬇️ Download Tailored Resume (DOCX)", data=docx_bytes, file_name="tailored_resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e:
                st.error(f"DOCX export failed: {e}. Ensure python-docx is in requirements.")
        elif orig_ext == ".pdf":
            try:
                pdf_bytes = export_pdf_text(tailored_resume)
                st.download_button("⬇️ Download Tailored Resume (PDF)", data=pdf_bytes, file_name="tailored_resume.pdf", mime="application/pdf")
            except Exception as e:
                st.error(f"PDF export failed: {e}. Ensure fpdf2 is in requirements.")
        else:
            # txt fallback
            st.download_button("⬇️ Download Tailored Resume (TXT)", data=tailored_resume, file_name="tailored_resume.txt", mime="text/plain")

        # Also offer cover letter downloads
        st.download_button("⬇️ Download Cover Letter (TXT)", data=cover_letter, file_name="cover_letter.txt", mime="text/plain")

st.caption("All processing is local to this Streamlit session. Downloads are text-based/ATS-friendly; they preserve content but not exact original visual styling.")
